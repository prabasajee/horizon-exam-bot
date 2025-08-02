"""
Flask API for Horizon Exam Bot
Handles PDF/DOCX uploads, text extraction, quiz creation, and quiz management
"""

import os
import json
import uuid
from datetime import datetime
from flask import Flask, request, jsonify, render_template
from werkzeug.utils import secure_filename
import PyPDF2
from docx import Document
from typing import Dict, List, Any
import re
from collections import Counter
import html
import secrets
from security_config import (
    SecurityConfig, InputValidator, SecurityLogger, 
    rate_limit, validate_input, security_logger
)

app = Flask(__name__)

# Security: Use environment variable for secret key or generate secure random key
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY') or os.urandom(24).hex()
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

# Security headers
@app.after_request
def set_security_headers(response):
    """Add security headers to all responses"""
    for header, value in SecurityConfig.SECURITY_HEADERS.items():
        response.headers[header] = value
    return response

# Ensure upload directory exists
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs('data', exist_ok=True)

# Store for quiz sessions (in production, use a database)
quiz_sessions = {}


class DocumentProcessor:
    """Handles text extraction from PDF and DOCX files"""
    
    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        """Extract text from PDF file"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        except Exception as e:
            raise Exception(f"Error reading PDF: {str(e)}")
        return text
    
    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """Extract text from DOCX file"""
        text = ""
        try:
            doc = Document(file_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
        except Exception as e:
            raise Exception(f"Error reading DOCX: {str(e)}")
        return text
    
    @classmethod
    def extract_text(cls, file_path: str, file_extension: str) -> str:
        """Extract text based on file extension"""
        if file_extension.lower() == '.pdf':
            return cls.extract_text_from_pdf(file_path)
        elif file_extension.lower() == '.docx':
            return cls.extract_text_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")


class NoteGenerator:
    """Simple note generator that creates study notes from text"""
    
    @staticmethod
    def generate_notes(text: str, style: str = "bullet") -> Dict[str, Any]:
        """Generate study notes from text"""
        try:
            # Clean text
            cleaned_text = NoteGenerator._clean_text(text)
            
            # Extract key information
            key_points = NoteGenerator._extract_key_points(cleaned_text)
            definitions = NoteGenerator._extract_definitions(cleaned_text)
            important_facts = NoteGenerator._extract_important_facts(cleaned_text)
            summary = NoteGenerator._create_summary(cleaned_text)
            
            # Format notes based on style
            if style == "bullet":
                formatted_notes = NoteGenerator._format_bullet_notes(key_points, definitions, important_facts)
            elif style == "numbered":
                formatted_notes = NoteGenerator._format_numbered_notes(key_points, definitions, important_facts)
            elif style == "paragraph":
                formatted_notes = NoteGenerator._format_paragraph_notes(summary, key_points)
            elif style == "flashcard":
                formatted_notes = NoteGenerator._format_flashcard_notes(definitions, important_facts)
            else:
                formatted_notes = NoteGenerator._format_bullet_notes(key_points, definitions, important_facts)
            
            return {
                'success': True,
                'notes': formatted_notes,
                'summary': summary,
                'key_points_count': len(key_points),
                'definitions_count': len(definitions),
                'facts_count': len(important_facts),
                'original_word_count': len(text.split()),
                'notes_word_count': len(formatted_notes.split()),
                'compression_ratio': round(len(formatted_notes.split()) / len(text.split()) * 100, 1) if text.split() else 0
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': str(e)
            }
    
    @staticmethod
    def _clean_text(text: str) -> str:
        """Clean and preprocess text"""
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text)
        # Remove special characters but keep basic punctuation
        text = re.sub(r'[^\w\s\.\,\!\?\;\:\-\(\)]', '', text)
        return text.strip()
    
    @staticmethod
    def _extract_key_points(text: str) -> List[str]:
        """Extract key points from text"""
        sentences = re.split(r'[.!?]+', text)
        key_points = []
        
        key_indicators = [
            'important', 'key', 'main', 'primary', 'essential', 'crucial',
            'significant', 'major', 'fundamental', 'basic', 'critical'
        ]
        
        for sentence in sentences:
            sentence = sentence.strip()
            if 20 <= len(sentence) <= 150:
                # Check for key indicators
                if any(indicator in sentence.lower() for indicator in key_indicators):
                    key_points.append(NoteGenerator._simplify_sentence(sentence))
                # Look for sentences with numbers
                elif re.search(r'\d+', sentence):
                    key_points.append(NoteGenerator._simplify_sentence(sentence))
        
        return key_points[:8]
    
    @staticmethod
    def _extract_definitions(text: str) -> List[Dict[str, str]]:
        """Extract definitions from text"""
        sentences = re.split(r'[.!?]+', text)
        definitions = []
        
        definition_patterns = [
            r'(.+?) is (.+?)$',
            r'(.+?) are (.+?)$',
            r'(.+?) means (.+?)$',
            r'(.+?) refers to (.+?)$'
        ]
        
        for sentence in sentences:
            if 20 <= len(sentence) <= 200:
                for pattern in definition_patterns:
                    match = re.search(pattern, sentence, re.IGNORECASE)
                    if match:
                        term = match.group(1).strip()
                        definition = match.group(2).strip()
                        
                        # Clean up term
                        term = re.sub(r'^(the|a|an)\s+', '', term.lower()).title()
                        definition = NoteGenerator._simplify_sentence(definition)
                        
                        if len(term) < 50 and len(definition) < 150:
                            definitions.append({
                                'term': term,
                                'definition': definition
                            })
                        break
        
        return definitions[:6]
    
    @staticmethod
    def _extract_important_facts(text: str) -> List[str]:
        """Extract important facts"""
        sentences = re.split(r'[.!?]+', text)
        facts = []
        
        for sentence in sentences:
            sentence = sentence.strip()
            if 15 <= len(sentence) <= 120:
                # Look for sentences with numbers
                if re.search(r'\d+', sentence):
                    facts.append(NoteGenerator._simplify_sentence(sentence))
                # Look for cause and effect
                elif any(word in sentence.lower() for word in ['because', 'due to', 'results in', 'causes']):
                    facts.append(NoteGenerator._simplify_sentence(sentence))
        
        return facts[:5]
    
    @staticmethod
    def _create_summary(text: str) -> str:
        """Create a brief summary"""
        sentences = re.split(r'[.!?]+', text)
        
        # Get first, middle, and last meaningful sentences
        meaningful_sentences = [s.strip() for s in sentences if 50 <= len(s.strip()) <= 200]
        
        if len(meaningful_sentences) >= 3:
            selected = [
                meaningful_sentences[0],
                meaningful_sentences[len(meaningful_sentences)//2],
                meaningful_sentences[-1]
            ]
        else:
            selected = meaningful_sentences[:3]
        
        return ' '.join([NoteGenerator._simplify_sentence(s) for s in selected])
    
    @staticmethod
    def _simplify_sentence(sentence: str) -> str:
        """Simplify sentence for easier reading"""
        replacements = {
            'utilize': 'use', 'demonstrate': 'show', 'implement': 'put in place',
            'facilitate': 'help', 'commence': 'start', 'terminate': 'end',
            'subsequently': 'then', 'furthermore': 'also', 'therefore': 'so',
            'approximately': 'about', 'sufficient': 'enough', 'numerous': 'many'
        }
        
        words = sentence.split()
        simplified_words = []
        
        for word in words:
            clean_word = re.sub(r'[^\w]', '', word.lower())
            if clean_word in replacements:
                # Preserve capitalization
                if word[0].isupper():
                    replacement = replacements[clean_word].capitalize()
                else:
                    replacement = replacements[clean_word]
                # Add back punctuation
                punct = ''.join([c for c in word if not c.isalnum()])
                simplified_words.append(replacement + punct)
            else:
                simplified_words.append(word)
        
        return ' '.join(simplified_words)
    
    @staticmethod
    def _format_bullet_notes(key_points: List[str], definitions: List[Dict], facts: List[str]) -> str:
        """Format as bullet points"""
        notes = "📝 **STUDY NOTES**\n\n"
        
        if key_points:
            notes += "🔑 **Key Points:**\n"
            for point in key_points:
                notes += f"• {point}\n"
            notes += "\n"
        
        if definitions:
            notes += "📚 **Important Terms:**\n"
            for def_item in definitions:
                notes += f"• **{def_item['term']}**: {def_item['definition']}\n"
            notes += "\n"
        
        if facts:
            notes += "💡 **Important Facts:**\n"
            for fact in facts:
                notes += f"• {fact}\n"
        
        return notes
    
    @staticmethod
    def _format_numbered_notes(key_points: List[str], definitions: List[Dict], facts: List[str]) -> str:
        """Format as numbered list"""
        notes = "📝 **STUDY NOTES**\n\n"
        counter = 1
        
        if key_points:
            notes += "🔑 **Key Points:**\n"
            for point in key_points:
                notes += f"{counter}. {point}\n"
                counter += 1
            notes += "\n"
        
        if definitions:
            notes += "📚 **Important Terms:**\n"
            for def_item in definitions:
                notes += f"{counter}. **{def_item['term']}**: {def_item['definition']}\n"
                counter += 1
            notes += "\n"
        
        if facts:
            notes += "💡 **Important Facts:**\n"
            for fact in facts:
                notes += f"{counter}. {fact}\n"
                counter += 1
        
        return notes
    
    @staticmethod
    def _format_paragraph_notes(summary: str, key_points: List[str]) -> str:
        """Format as paragraphs"""
        notes = "📝 **STUDY NOTES**\n\n"
        notes += "📖 **Summary:**\n"
        notes += f"{summary}\n\n"
        
        if key_points:
            notes += "🔑 **Main Ideas:**\n"
            notes += "The key concepts to remember are: " + ", ".join(key_points[:5]) + ".\n\n"
        
        return notes
    
    @staticmethod
    def _format_flashcard_notes(definitions: List[Dict], facts: List[str]) -> str:
        """Format as flashcards"""
        notes = "🃏 **FLASHCARD NOTES**\n\n"
        
        if definitions:
            notes += "📚 **Term Cards:**\n"
            for i, def_item in enumerate(definitions, 1):
                notes += f"**Card {i}:**\n"
                notes += f"Front: What is {def_item['term']}?\n"
                notes += f"Back: {def_item['definition']}\n\n"
        
        if facts:
            notes += "💡 **Fact Cards:**\n"
            for i, fact in enumerate(facts, len(definitions) + 1):
                notes += f"**Card {i}:**\n"
                notes += f"Fact: {fact}\n\n"
        
        return notes


class QuizManager:
    """Manages quiz creation, storage, and scoring"""
    
    @staticmethod
    def save_quiz(quiz_data: Dict[str, Any]) -> str:
        """Save quiz to JSON file and return quiz ID"""
        quiz_id = str(uuid.uuid4())
        quiz_data['id'] = quiz_id
        quiz_data['created_at'] = datetime.now().isoformat()
        
        filename = f"data/quiz_{quiz_id}.json"
        with open(filename, 'w', encoding='utf-8') as f:
            json.dump(quiz_data, f, indent=2, ensure_ascii=False)
        
        return quiz_id
    
    @staticmethod
    def load_quiz(quiz_id: str) -> Dict[str, Any]:
        """Load quiz from JSON file"""
        filename = f"data/quiz_{quiz_id}.json"
        if not os.path.exists(filename):
            raise FileNotFoundError(f"Quiz {quiz_id} not found")
        
        with open(filename, 'r', encoding='utf-8') as f:
            return json.load(f)
    
    @staticmethod
    def calculate_score(quiz_data: Dict[str, Any], user_answers: Dict[str, str]) -> Dict[str, Any]:
        """Calculate quiz score and provide feedback"""
        total_questions = len(quiz_data['questions'])
        correct_answers = 0
        detailed_results = []
        
        for i, question in enumerate(quiz_data['questions']):
            question_num = str(i)
            user_answer = user_answers.get(question_num, "")
            correct_answer = question['correct_answer']
            is_correct = user_answer.strip().lower() == correct_answer.strip().lower()
            
            if is_correct:
                correct_answers += 1
            
            detailed_results.append({
                'question': question['question'],
                'your_answer': user_answer,
                'correct_answer': correct_answer,
                'is_correct': is_correct,
                'explanation': question.get('explanation', '')
            })
        
        score_percentage = (correct_answers / total_questions) * 100 if total_questions > 0 else 0
        
        return {
            'total_questions': total_questions,
            'correct_answers': correct_answers,
            'score_percentage': round(score_percentage, 2),
            'detailed_results': detailed_results,
            'timestamp': datetime.now().isoformat()
        }


def allowed_file(filename: str) -> bool:
    """Check if file extension is allowed"""
    if not filename:
        return False
    
    # Use InputValidator for comprehensive validation
    validation_result = InputValidator.validate_filename(filename)
    return validation_result['valid']

def validate_file_content(file_path: str, expected_extension: str) -> bool:
    """Validate file content matches expected type"""
    try:
        if expected_extension.lower() == '.pdf':
            # Try to read PDF to validate it's actually a PDF
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                # Just try to access pages to validate format
                len(pdf_reader.pages)
        elif expected_extension.lower() == '.docx':
            # Try to read DOCX to validate format
            Document(file_path)
        return True
    except Exception:
        return False


@app.route('/')
def index():
    """Main page with upload form"""
    return render_template('index.html')


@app.route('/api/upload', methods=['POST'])
@rate_limit(limit=10, window=3600)  # 10 uploads per hour
def upload_file():
    """Handle file upload and text extraction"""
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        # Validate filename
        filename_validation = InputValidator.validate_filename(file.filename)
        if not filename_validation['valid']:
            security_logger.log_validation_error(
                'File upload',
                f"Invalid filename: {filename_validation['error']}",
                request.remote_addr
            )
            return jsonify({'error': filename_validation['error']}), 400
        
        # Additional security checks
        original_filename = file.filename
        
        # Generate secure filename
        filename = secure_filename(original_filename)
        if not filename:  # secure_filename might return empty string for malicious filenames
            filename = f"upload_{secrets.token_hex(8)}.{original_filename.rsplit('.', 1)[1].lower()}"
        
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        
        # Check file size before saving
        file.seek(0, 2)  # Seek to end
        file_size = file.tell()
        file.seek(0)  # Reset to beginning
        
        if file_size > SecurityConfig.MAX_FILE_SIZE:
            return jsonify({'error': 'File too large. Maximum size is 16MB.'}), 413
        
        if file_size == 0:
            return jsonify({'error': 'Empty file not allowed'}), 400
        
        # Log file upload
        security_logger.log_file_upload(original_filename, file_size, request.remote_addr)
        
        # Save uploaded file
        file.save(file_path)
        
        # Validate file content
        file_extension = os.path.splitext(filename)[1]
        if not validate_file_content(file_path, file_extension):
            os.remove(file_path)  # Clean up invalid file
            security_logger.log_validation_error(
                'File content validation',
                f"Invalid file format for {original_filename}",
                request.remote_addr
            )
            return jsonify({'error': 'Invalid file format or corrupted file'}), 400
        
        # Extract text
        extracted_text = DocumentProcessor.extract_text(file_path, file_extension)
        
        # Sanitize extracted text
        extracted_text = InputValidator.sanitize_text(extracted_text, max_length=SecurityConfig.MAX_TEXT_LENGTH)
        
        # Clean up uploaded file
        os.remove(file_path)
        
        return jsonify({
            'success': True,
            'filename': InputValidator.sanitize_text(original_filename, max_length=SecurityConfig.MAX_FILENAME_LENGTH),
            'extracted_text': extracted_text,
            'text_length': len(extracted_text)
        })
    
    except Exception as e:
        # Ensure uploaded file is cleaned up on error
        if 'file_path' in locals() and os.path.exists(file_path):
            try:
                os.remove(file_path)
            except:
                pass
        
        # Log error securely (don't expose internal details)
        app.logger.error(f"Upload error: {str(e)}")
        return jsonify({'error': 'Server error occurred during file processing'}), 500


@app.route('/api/quiz/create', methods=['POST'])
@rate_limit(limit=20, window=3600)  # 20 quiz creations per hour
@validate_input({'required_fields': ['title', 'questions'], 'max_fields': 10})
def create_quiz():
    """Create a new quiz from questions and answers"""
    try:
        data = request.get_json()
        
        # Validate quiz data structure and content
        validation_result = InputValidator.validate_quiz_data(data)
        if not validation_result['valid']:
            security_logger.log_validation_error(
                'Quiz creation',
                f"Validation errors: {validation_result['errors']}",
                request.remote_addr
            )
            return jsonify({'error': validation_result['errors'][0]}), 400
        
        # Sanitize title and description
        title = InputValidator.sanitize_text(data['title'], max_length=SecurityConfig.MAX_TITLE_LENGTH)
        description = InputValidator.sanitize_text(data.get('description', ''), max_length=SecurityConfig.MAX_DESCRIPTION_LENGTH)
        
        # Sanitize questions
        sanitized_questions = []
        for question in data['questions']:
            sanitized_question = {
                'question': InputValidator.sanitize_text(question['question'], max_length=SecurityConfig.MAX_QUESTION_LENGTH),
                'correct_answer': InputValidator.sanitize_text(question['correct_answer'], max_length=SecurityConfig.MAX_OPTION_LENGTH),
                'explanation': InputValidator.sanitize_text(question.get('explanation', ''), max_length=SecurityConfig.MAX_DESCRIPTION_LENGTH)
            }
            
            # Sanitize options
            sanitized_options = []
            for option in question['options']:
                sanitized_option = InputValidator.sanitize_text(str(option), max_length=SecurityConfig.MAX_OPTION_LENGTH)
                if sanitized_option:
                    sanitized_options.append(sanitized_option)
            
            sanitized_question['options'] = sanitized_options
            sanitized_questions.append(sanitized_question)
        
        # Create sanitized quiz data
        quiz_data = {
            'title': title,
            'description': description,
            'questions': sanitized_questions
        }
        
        # Save quiz
        quiz_id = QuizManager.save_quiz(quiz_data)
        
        return jsonify({
            'success': True,
            'quiz_id': quiz_id,
            'message': 'Quiz created successfully'
        })
    
    except Exception as e:
        app.logger.error(f"Quiz creation error: {str(e)}")
        return jsonify({'error': 'Failed to create quiz'}), 500


@app.route('/api/quiz/<quiz_id>', methods=['GET'])
def get_quiz(quiz_id: str):
    """Get quiz questions for taking the quiz"""
    try:
        quiz_data = QuizManager.load_quiz(quiz_id)
        
        # Return quiz without correct answers
        quiz_for_user = {
            'id': quiz_data['id'],
            'title': quiz_data['title'],
            'description': quiz_data.get('description', ''),
            'questions': []
        }
        
        for question in quiz_data['questions']:
            quiz_for_user['questions'].append({
                'question': question['question'],
                'options': question['options']
            })
        
        return jsonify(quiz_for_user)
    
    except FileNotFoundError:
        return jsonify({'error': 'Resource not found'}), 404
    except Exception as e:
        app.logger.error(f"Quiz retrieval error: {str(e)}")
        return jsonify({'error': 'Server error occurred'}), 500


@app.route('/api/quiz/<quiz_id>/submit', methods=['POST'])
def submit_quiz(quiz_id: str):
    """Submit quiz answers and get results"""
    try:
        quiz_data = QuizManager.load_quiz(quiz_id)
        user_answers = request.get_json()
        
        if not user_answers or 'answers' not in user_answers:
            return jsonify({'error': 'No answers provided'}), 400
        
        # Calculate score and feedback
        results = QuizManager.calculate_score(quiz_data, user_answers['answers'])
        
        # Store session results
        session_id = str(uuid.uuid4())
        quiz_sessions[session_id] = {
            'quiz_id': quiz_id,
            'results': results,
            'user_info': user_answers.get('user_info', {})
        }
        
        return jsonify({
            'success': True,
            'session_id': session_id,
            'results': results
        })
    
    except FileNotFoundError:
        return jsonify({'error': 'Resource not found'}), 404
    except Exception as e:
        app.logger.error(f"Quiz submission error: {str(e)}")
        return jsonify({'error': 'Server error occurred'}), 500


@app.route('/api/quiz/list', methods=['GET'])
def list_quizzes():
    """List all available quizzes"""
    try:
        quizzes = []
        data_dir = 'data'
        
        if os.path.exists(data_dir):
            for filename in os.listdir(data_dir):
                if filename.startswith('quiz_') and filename.endswith('.json'):
                    quiz_id = filename[5:-5]  # Remove 'quiz_' prefix and '.json' suffix
                    try:
                        quiz_data = QuizManager.load_quiz(quiz_id)
                        quizzes.append({
                            'id': quiz_data['id'],
                            'title': quiz_data['title'],
                            'description': quiz_data.get('description', ''),
                            'question_count': len(quiz_data['questions']),
                            'created_at': quiz_data.get('created_at', '')
                        })
                    except Exception:
                        continue  # Skip invalid quiz files
        
        return jsonify({'quizzes': quizzes})
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/session/<session_id>/results', methods=['GET'])
def get_session_results(session_id: str):
    """Get results for a specific quiz session"""
    if session_id not in quiz_sessions:
        return jsonify({'error': 'Session not found'}), 404
    
    return jsonify(quiz_sessions[session_id])


@app.route('/quiz/<quiz_id>')
def quiz_page(quiz_id: str):
    """Render quiz taking page"""
    return render_template('quiz.html', quiz_id=quiz_id)


@app.route('/results/<session_id>')
def results_page(session_id: str):
    """Render quiz results page"""
    return render_template('results.html', session_id=session_id)


@app.route('/notes')
def notes_page():
    """Render notes listing page"""
    return render_template('notes.html')


@app.errorhandler(413)
def file_too_large(error):
    return jsonify({'error': 'File too large. Maximum size is 16MB.'}), 413


@app.errorhandler(500)
def internal_error(error):
    return jsonify({'error': 'Internal server error. Please try again.'}), 500


@app.errorhandler(404)
def not_found(error):
    return jsonify({'error': 'Endpoint not found.'}), 404


@app.route('/api/notes/generate', methods=['POST'])
@rate_limit(limit=30, window=3600)  # 30 note generations per hour
def generate_notes():
    """Generate study notes from text"""
    try:
        data = request.get_json()
        
        if not data:
            return jsonify({'success': False, 'error': 'No data provided'})
        
        # Validate and sanitize input
        text = InputValidator.sanitize_text(data.get('text', ''), max_length=SecurityConfig.MAX_TEXT_LENGTH)
        if not text:
            return jsonify({'success': False, 'error': 'No valid text provided'})
        
        if len(text) < SecurityConfig.MIN_TEXT_LENGTH_FOR_NOTES:
            return jsonify({'success': False, 'error': f'Text too short (minimum {SecurityConfig.MIN_TEXT_LENGTH_FOR_NOTES} characters)'})
        
        # Validate style parameter
        allowed_styles = {'bullet', 'numbered', 'paragraph', 'flashcard'}
        style = data.get('style', 'bullet')
        if style not in allowed_styles:
            style = 'bullet'
        
        result = NoteGenerator.generate_notes(text, style)
        return jsonify(result)
        
    except Exception as e:
        app.logger.error(f"Notes generation error: {str(e)}")
        return jsonify({'success': False, 'error': 'Failed to generate notes'}), 500


@app.route('/api/notes/save', methods=['POST'])
@rate_limit(limit=50, window=3600)  # 50 note saves per hour
def save_notes():
    """Save generated notes"""
    try:
        data = request.get_json()
        
        if not data:
            return jsonify({'success': False, 'error': 'No data provided'})
        
        # Validate and sanitize input
        title = InputValidator.sanitize_text(data.get('title', ''), max_length=SecurityConfig.MAX_TITLE_LENGTH)
        content = InputValidator.sanitize_text(data.get('content', ''), max_length=SecurityConfig.MAX_TEXT_LENGTH)
        source_content = InputValidator.sanitize_text(data.get('source_content', ''), max_length=SecurityConfig.MAX_TEXT_LENGTH)
        
        if not title or not content:
            return jsonify({'success': False, 'error': 'Title and content are required'})
        
        # Save notes to file
        notes_id = str(uuid.uuid4())
        notes_data = {
            'id': notes_id,
            'title': title,
            'content': content,
            'source_content': source_content,
            'created_at': datetime.now().isoformat()
        }
        
        filename = f"data/notes_{notes_id}.json"
        with open(filename, 'w', encoding='utf-8') as f:
            json.dump(notes_data, f, indent=2, ensure_ascii=False)
        
        return jsonify({
            'success': True,
            'notes_id': notes_id,
            'message': 'Notes saved successfully'
        })
        
    except Exception as e:
        app.logger.error(f"Notes save error: {str(e)}")
        return jsonify({'success': False, 'error': 'Failed to save notes'}), 500


@app.route('/api/notes/list', methods=['GET'])
def list_notes():
    """List all saved notes"""
    try:
        notes = []
        data_dir = 'data'
        
        if os.path.exists(data_dir):
            for filename in os.listdir(data_dir):
                if filename.startswith('notes_') and filename.endswith('.json'):
                    try:
                        with open(os.path.join(data_dir, filename), 'r', encoding='utf-8') as f:
                            notes_data = json.load(f)
                            notes.append({
                                'id': notes_data['id'],
                                'title': notes_data['title'],
                                'created_at': notes_data.get('created_at', ''),
                                'preview': notes_data['content'][:100] + '...' if len(notes_data['content']) > 100 else notes_data['content']
                            })
                    except Exception:
                        continue
        
        return jsonify({'success': True, 'notes': notes})
        
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/notes/styles', methods=['GET'])
def get_note_styles():
    """Get available note styles"""
    styles = [
        {
            'id': 'bullet',
            'name': 'Bullet Points',
            'description': 'Organized bullet points with key topics',
            'icon': 'fas fa-list-ul'
        },
        {
            'id': 'numbered',
            'name': 'Numbered List',
            'description': 'Sequential numbered points for easy reference',
            'icon': 'fas fa-list-ol'
        },
        {
            'id': 'paragraph',
            'name': 'Summary Paragraphs',
            'description': 'Flowing summary with main ideas',
            'icon': 'fas fa-align-left'
        },
        {
            'id': 'flashcard',
            'name': 'Flashcards',
            'description': 'Question and answer format for studying',
            'icon': 'fas fa-layer-group'
        }
    ]
    
    return jsonify({
        'success': True,
        'styles': styles
    })


if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)
